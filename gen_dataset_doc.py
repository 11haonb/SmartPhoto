"""生成 SmartPhoto 测试数据集采集任务书 Word 文档"""
import docx
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 样式设置 =====
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.5

# ===== 标题 =====
title = doc.add_heading('SmartPhoto 测试数据集采集任务书', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===== 项目说明 =====
doc.add_heading('一、项目背景', level=1)
doc.add_paragraph(
    'SmartPhoto 是一款 AI 智能照片整理应用，具备以下核心功能：'
)
features = [
    'EXIF 时间线提取 — 自动读取拍摄时间、相机型号、GPS 位置，按日期分组',
    '质量检测 — 识别模糊、过曝、欠曝、截图等问题照片',
    '智能分类 — 将照片自动归类为人物/风景/美食/文档/截图/其他',
    '相似分组 — 通过感知哈希（pHash）找出重复和相似照片',
    '最佳挑选 — 从每组相似照片中自动选出质量最好的一张',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph(
    '为验证以上所有功能，需要采集一批覆盖各种场景的测试照片。'
    '本文档列出详细的采集清单和拍摄要求。'
)

# ===== 总体要求 =====
doc.add_heading('二、总体要求', level=1)

reqs = [
    ('格式', 'JPG 或 PNG，不要 HEIF/HEIC'),
    ('单张大小', '不超过 10MB'),
    ('拍摄设备', '手机拍摄，保留原始 EXIF 信息'),
    ('传输方式', '通过数据线、AirDrop 或网盘「原图」导出。严禁通过微信/QQ 传图（会压缩并丢失 EXIF）'),
    ('拍摄周期', '分 3 天拍摄（或修改手机日期模拟，确保照片分布在 3 个不同日期）'),
    ('GPS 定位', '拍摄时确保手机开启了定位服务（至少 5 张需要有 GPS 信息）'),
    ('命名规则', '按编号命名，如 01-风景-公园.jpg、15-模糊-晃动.jpg'),
    ('交付方式', '全部照片放入一个文件夹，打包成 zip 压缩包交付'),
]

table = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = '项目'
hdr[1].text = '要求'
for item, req in reqs:
    row = table.add_row().cells
    row[0].text = item
    row[1].text = req

# ===== 数量汇总 =====
doc.add_heading('三、数量汇总', level=1)
doc.add_paragraph('总计需要采集约 45-50 张照片，分为 6 个大类：')

summary_table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = summary_table.rows[0].cells
hdr[0].text = '类别'
hdr[1].text = '数量'
hdr[2].text = '测试功能'
summary_data = [
    ('正常高质量照片', '16 张', '时间线分组、智能分类、GPS 提取'),
    ('问题照片', '10 张', '质量检测（模糊/过曝/欠曝/截图）'),
    ('文档类照片', '4 张', '文档分类识别'),
    ('连拍相似组', '12 张（4组×3张）', '相似分组 + 最佳挑选'),
    ('无 EXIF 图片', '2 张', 'EXIF 缺失处理'),
    ('特殊格式', '2 张', 'PNG 格式支持'),
]
for cat, count, func in summary_data:
    row = summary_table.add_row().cells
    row[0].text = cat
    row[1].text = count
    row[2].text = func

# ===== 详细采集清单 =====
doc.add_heading('四、详细采集清单', level=1)

# --- 4.1 人物类 ---
doc.add_heading('4.1 人物类（person）— 6 张', level=2)
doc.add_paragraph('测试人物分类及子类别识别（portrait/group/selfie）。')

data_person = [
    ('01', '自拍-正面', '第 1 天', '前置摄像头自拍，露出面部，光线充足', '开', 'selfie'),
    ('02', '自拍-户外', '第 1 天', '户外环境自拍，有背景', '开', 'selfie'),
    ('03', '人像-半身', '第 2 天', '他人拍摄的半身照，背景简洁', '开', 'portrait'),
    ('04', '人像-全身', '第 2 天', '全身站立照，清晰对焦', '可选', 'portrait'),
    ('05', '合照-2人', '第 3 天', '两人合照', '开', 'group'),
    ('06', '合照-多人', '第 3 天', '3 人以上合照或集体照', '开', 'group'),
]

table = doc.add_table(rows=1, cols=6, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['编号', '内容', '拍摄日期', '拍摄要求', 'GPS', '子类别']):
    table.rows[0].cells[i].text = h
for row_data in data_person:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# --- 4.2 风景类 ---
doc.add_heading('4.2 风景类（landscape）— 6 张', level=2)
doc.add_paragraph('测试风景分类及子类别识别（nature/building/city）。')

data_landscape = [
    ('07', '自然风光-山/水/树', '第 1 天', '公园、湖边、山景等自然场景', '开', 'nature'),
    ('08', '自然风光-天空/云', '第 1 天', '蓝天白云或日落', '开', 'nature'),
    ('09', '建筑-单栋', '第 2 天', '一栋有特色的建筑物', '开', 'building'),
    ('10', '建筑-室内', '第 2 天', '室内空间（大厅、走廊等）', '可选', 'building'),
    ('11', '城市街景-白天', '第 3 天', '街道、车流、商铺', '开', 'city'),
    ('12', '城市街景-夜景', '第 3 天', '夜晚的城市灯光', '可选', 'city'),
]

table = doc.add_table(rows=1, cols=6, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['编号', '内容', '拍摄日期', '拍摄要求', 'GPS', '子类别']):
    table.rows[0].cells[i].text = h
for row_data in data_landscape:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# --- 4.3 美食类 ---
doc.add_heading('4.3 美食类（food）— 4 张', level=2)
doc.add_paragraph('测试美食分类识别。')

data_food = [
    ('13', '中餐菜品', '第 1 天', '一盘菜的俯拍或侧拍，画面清晰', '可选'),
    ('14', '饮品/甜点', '第 2 天', '咖啡、奶茶、蛋糕等', '可选'),
    ('15', '多道菜全景', '第 3 天', '一桌菜的全景照', '可选'),
    ('16', '水果/食材', '第 3 天', '水果或未加工食材的特写', '可选'),
]

table = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['编号', '内容', '拍摄日期', '拍摄要求', 'GPS']):
    table.rows[0].cells[i].text = h
for row_data in data_food:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# --- 4.4 文档类 ---
doc.add_heading('4.4 文档类（document）— 4 张', level=2)
doc.add_paragraph('测试文档分类识别。')

data_doc = [
    ('17', '书本页面', '任意', '拍一页书的文字内容，尽量拍正拍清晰'),
    ('18', '手写笔记', '任意', '拍一页手写的笔记或草稿'),
    ('19', '白板/黑板', '任意', '拍会议白板或教室黑板上的内容'),
    ('20', '电脑屏幕/PPT', '任意', '拍显示器上的 PPT 或表格'),
]

table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['编号', '内容', '拍摄日期', '拍摄要求']):
    table.rows[0].cells[i].text = h
for row_data in data_doc:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# --- 4.5 问题照片 ---
doc.add_heading('4.5 问题照片 — 10 张', level=2)
doc.add_paragraph('测试质量检测功能，需要故意拍出有问题的照片。')

data_issues = [
    ('21', '模糊-运动', '任意', '拍照时故意快速晃动手机，产生运动模糊'),
    ('22', '模糊-失焦', '任意', '手动对焦到最近处，让主体完全虚化'),
    ('23', '模糊-手抖', '任意', '单手拍照故意手抖，产生轻微模糊'),
    ('24', '过曝-强光', '任意', '对着窗户/太阳/灯光拍，画面大面积发白'),
    ('25', '过曝-反光', '任意', '对着反光的白色桌面或镜面拍'),
    ('26', '欠曝-太暗', '任意', '在暗处拍照不开闪光灯，画面发黑'),
    ('27', '欠曝-遮挡', '任意', '用手指遮住大半镜头拍一张'),
    ('28', '截图-主屏幕', '任意', '手机截屏：主屏幕（带顶部状态栏+底部导航）'),
    ('29', '截图-聊天', '任意', '手机截屏：微信/短信聊天记录页面'),
    ('30', '截图-设置', '任意', '手机截屏：系统设置页面或 App 内页面'),
]

table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['编号', '内容', '拍摄日期', '拍摄要求']):
    table.rows[0].cells[i].text = h
for row_data in data_issues:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

p = doc.add_paragraph()
run = p.add_run('提示：')
run.bold = True
p.add_run('模糊照片要确保肉眼可见模糊，不要只是轻微抖动。过曝要画面大片发白，欠曝要画面大片发黑。')

# --- 4.6 连拍相似组 ---
doc.add_heading('4.6 连拍相似组 — 12 张（4 组，每组 3 张）', level=2)
doc.add_paragraph(
    '测试 pHash 相似分组和最佳挑选功能。每组对同一个场景连续拍 3 张，'
    '其中 1 张拍好、2 张故意拍差。'
)

doc.add_heading('组 A：桌面静物', level=3)
data_group_a = [
    ('31', '桌面静物-最佳', '任意', '对着桌上的杯子/花/摆件，认真构图，对焦清晰，这张应该是最好的'),
    ('32', '桌面静物-略差', '任意', '同一物体，不移动，稍微歪一点或手抖一下拍'),
    ('33', '桌面静物-最差', '任意', '同一物体，不移动，故意晃一下或对焦偏移'),
]

table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['编号', '内容', '拍摄日期', '拍摄要求']):
    table.rows[0].cells[i].text = h
for row_data in data_group_a:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_heading('组 B：窗外/阳台景色', level=3)
data_group_b = [
    ('34', '窗外景色-最佳', '任意', '对着窗外景色，构图稳定，对焦清晰'),
    ('35', '窗外景色-略差', '任意', '同一角度，稍微移动或轻微手抖'),
    ('36', '窗外景色-最差', '任意', '同一角度，故意晃动'),
]

table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['编号', '内容', '拍摄日期', '拍摄要求']):
    table.rows[0].cells[i].text = h
for row_data in data_group_b:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_heading('组 C：人物同一姿势', level=3)
data_group_c = [
    ('37', '人物姿势-最佳', '任意', '请一个人站好不动，拍一张清晰的'),
    ('38', '人物姿势-略差', '任意', '同一人不动，稍微偏移构图再拍'),
    ('39', '人物姿势-最差', '任意', '同一人不动，故意手抖或失焦拍'),
]

table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['编号', '内容', '拍摄日期', '拍摄要求']):
    table.rows[0].cells[i].text = h
for row_data in data_group_c:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_heading('组 D：美食同一盘菜', level=3)
data_group_d = [
    ('40', '同一盘菜-最佳', '任意', '对着一盘菜俯拍，光线好，对焦清晰'),
    ('41', '同一盘菜-略差', '任意', '同一盘菜不移动，换个角度或轻微手抖'),
    ('42', '同一盘菜-最差', '任意', '同一盘菜不移动，故意拍歪或虚焦'),
]

table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['编号', '内容', '拍摄日期', '拍摄要求']):
    table.rows[0].cells[i].text = h
for row_data in data_group_d:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

p = doc.add_paragraph()
run = p.add_run('关键要求：')
run.bold = True
p.add_run('每组 3 张照片必须是同一个场景/物体，拍摄时不要移动被拍物体，只改变拍摄质量。'
           '这样 pHash 值才会接近，系统才能正确分组。')

# --- 4.7 特殊测试 ---
doc.add_heading('4.7 特殊测试 — 4 张', level=2)

data_special = [
    ('43', 'PNG 格式照片', '任意', '将一张手机照片转为 PNG 格式保存（或直接截图为 PNG）'),
    ('44', 'PNG 格式照片 2', '任意', '另一张 PNG 格式的照片'),
    ('45', '网络下载图片', '任意', '从网上下载一张图片（无 EXIF），测试 EXIF 缺失处理'),
    ('46', '动物/宠物', '任意', '拍宠物或路边的猫/狗/鸟，测试 other 分类'),
]

table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['编号', '内容', '拍摄日期', '拍摄要求']):
    table.rows[0].cells[i].text = h
for row_data in data_special:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# ===== 验收标准 =====
doc.add_heading('五、验收标准', level=1)

checks = [
    '总数：46 张照片',
    '格式：全部为 JPG 或 PNG',
    '命名：按编号命名（01-xxx.jpg ~ 46-xxx.jpg）',
    '日期分布：至少分布在 3 个不同日期',
    'GPS 信息：至少 5 张包含 GPS 定位信息',
    '问题照片：模糊/过曝/欠曝肉眼可明显辨别',
    '连拍组：每组 3 张场景一致、质量有梯度差异',
    '传输：原图传输，未经微信/QQ 压缩',
]

for c in checks:
    doc.add_paragraph(c, style='List Bullet')

# ===== 文件命名参考 =====
doc.add_heading('六、文件命名参考', level=1)

naming = [
    '01-自拍-正面.jpg',
    '02-自拍-户外.jpg',
    '...',
    '07-风景-公园.jpg',
    '...',
    '13-美食-中餐.jpg',
    '...',
    '17-文档-书本.jpg',
    '...',
    '21-模糊-运动.jpg',
    '22-模糊-失焦.jpg',
    '...',
    '28-截图-主屏幕.png',
    '...',
    '31-连拍A-最佳.jpg',
    '32-连拍A-略差.jpg',
    '33-连拍A-最差.jpg',
    '...',
    '45-网络下载图.jpg',
    '46-动物-猫.jpg',
]

p = doc.add_paragraph()
for n in naming:
    p.add_run(n + '\n')

# ===== 注意事项 =====
doc.add_heading('七、注意事项', level=1)

notes = [
    ('严禁微信传图', '微信会自动压缩照片并删除 EXIF 信息（拍摄时间、GPS 等），导致测试数据无效。请使用数据线、AirDrop、或网盘「原图下载」功能传输。'),
    ('连拍不要移动物体', '连拍组的 3 张照片必须是同一场景，物体位置不变，只改变拍摄质量（抖/歪/虚焦）。移动物体会导致 pHash 差异过大，无法测试相似分组功能。'),
    ('过曝/欠曝要明显', '轻微的亮度差异不会被检测到。过曝照片应该有大片纯白区域，欠曝照片应该有大片纯黑区域。'),
    ('截图要完整', '截图需要包含手机顶部状态栏（时间、电量）和底部导航栏，这是截图识别的关键特征。'),
    ('确认 GPS 开启', '拍摄前到手机 设置 > 隐私 > 定位服务 确认相机 App 有定位权限。拍完一张后在相册里查看照片详情，确认有位置信息。'),
]

for title, desc in notes:
    p = doc.add_paragraph()
    run = p.add_run(title + '：')
    run.bold = True
    p.add_run(desc)

# ===== 保存 =====
output_path = '/home/liujinqi/smartphoto/SmartPhoto测试数据集采集任务书.docx'
doc.save(output_path)
print(f'文档已保存到: {output_path}')
